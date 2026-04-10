"""PDF → Word 转换后端（普通模式 / OCR 模式）+ 辅助函数"""


def convert_to_word(pdf_path, out_path, page_range=None, out_opts=None):
    """
    普通模式：直接提取文字转换（适合文字版PDF）。
    若 pdf2docx 失败则自动降级用 pypdf 提取纯文本写入 docx。
    """
    opts = out_opts or {}
    kwargs = dict(
        multi_sections    = opts.get("detect_toc", True),
        detect_table      = True,
        detect_page_break = True,
    )
    # ── 主方案：pdf2docx ──
    try:
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        if page_range:
            cv.convert(out_path, pages=page_range, **kwargs)
        else:
            cv.convert(out_path, start=0, end=None, **kwargs)
        cv.close()
        return
    except Exception:
        pass   # 进入降级方案

    # ── 降级方案：pypdf 提取纯文本 → python-docx 写入 ──
    from pypdf import PdfReader
    from docx import Document
    from docx.shared import Pt, Cm

    reader = PdfReader(pdf_path)
    pages  = reader.pages
    if page_range:
        pages = [pages[i] for i in page_range if i < len(pages)]

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    for i, page in enumerate(pages):
        text = page.extract_text() or ""
        for line in text.splitlines():
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(10.5)
        if i < len(pages) - 1:
            doc.add_page_break()

    doc.save(out_path)


def convert_to_word_ocr(pdf_path, out_path, page_range=None,
                        lang="chi_sim+eng", progress_cb=None):
    """
    OCR模式：将每页渲染为图片 → easyocr识别文字 → 写入Word
    适合扫描版PDF（内容为图片，无法直接提取文字）
    """
    import easyocr
    import numpy as np
    from pdf2image import convert_from_path
    from docx import Document
    from docx.shared import Pt, Cm

    lang_map = {
        "chi_sim+eng": ["ch_sim", "en"],
        "eng":         ["en"],
        "chi_tra+eng": ["ch_tra", "en"],
        "jpn+eng":     ["ja",    "en"],
    }
    ocr_langs = lang_map.get(lang, ["ch_sim", "en"])
    reader = easyocr.Reader(ocr_langs, gpu=False)

    images = convert_from_path(pdf_path, dpi=200)
    if page_range:
        images = [images[i] for i in page_range if i < len(images)]

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    total = len(images)
    for i, img in enumerate(images):
        img_np = np.array(img)
        results = reader.readtext(img_np, detail=0, paragraph=True)
        for line in results:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(10.5)
        if i < total - 1:
            doc.add_page_break()
        if progress_cb:
            progress_cb(int((i + 1) / total * 100))

    doc.save(out_path)


def detect_pdf_type(pdf_path, sample_pages=3):
    """
    检测 PDF 是否为扫描版（图片型）还是文字版。
    每页平均字符数 < 50 则判定为扫描版。
    返回 "scanned" 或 "text"
    """
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        total  = len(reader.pages)
        pages  = reader.pages[:min(sample_pages, total)]
        chars  = sum(len((p.extract_text() or "").strip()) for p in pages)
        avg    = chars / max(len(pages), 1)
        return "scanned" if avg < 50 else "text"
    except Exception:
        return "text"


def parse_page_range(text, total_pages):
    """解析 '1,3-5,8' 格式页码为0起始列表"""
    pages = set()
    for part in text.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            try:
                pages.update(range(int(a) - 1, int(b)))
            except ValueError:
                raise ValueError(f"页码格式错误: {part}")
        else:
            try:
                pages.add(int(part) - 1)
            except ValueError:
                raise ValueError(f"页码格式错误: {part}")
    valid = [p for p in sorted(pages) if 0 <= p < total_pages]
    if not valid:
        raise ValueError("所选页码超出文件范围")
    return valid
